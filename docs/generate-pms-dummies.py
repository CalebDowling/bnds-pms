"""Generate updated PMS_for_Dummies.docx"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
green = RGBColor(0x40, 0x72, 0x1D)
gray = RGBColor(0x6B, 0x6B, 0x6B)
light = RGBColor(0x99, 0x99, 0x99)

def green_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = green
    return h

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("The BNDS PMS\nfor Dummies")
run.font.size = Pt(28)
run.font.color.rgb = green
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A plain-English guide to everything behind the scenes")
run.font.size = Pt(12)
run.font.color.rgb = gray

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Boudreaux's Compounding Pharmacy\nUpdated April 2026")
run.font.size = Pt(10)
run.font.color.rgb = light

doc.add_page_break()

# What Is a PMS
green_heading("What Is a PMS?")
doc.add_paragraph("PMS stands for Pharmacy Management System. Think of it as the brain of the pharmacy. It's the software that keeps track of every prescription, every patient, every pill in the bottle, and every insurance claim. Without it, the pharmacy couldn't function.")
doc.add_paragraph("Our PMS (called BNDS PMS) is a custom system we built from scratch to replace our old one called DRX. We built our own because we needed features that off-the-shelf systems don't offer \u2014 especially around compounding, delivery management, and clinic billing.")

# Big Picture
green_heading("The Big Picture: How It All Connects")
doc.add_paragraph("Our PMS doesn't work alone. It talks to about a dozen other services and systems. Here's the simplest way to think about the prescription flow:")
for s in [
    "A doctor sends a prescription (SureScripts delivers it electronically)",
    "Our PMS receives it and a tech reviews it in the Intake queue",
    "We check if insurance covers it (Change Healthcare routes this to the insurance company)",
    "The label gets printed and the tech fills the prescription",
    "The tech scans the drug barcode to verify it's the right medication (patient safety)",
    "A pharmacist verifies everything is correct and signs off",
    "The prescription goes to the Waiting Bin for pickup, or gets shipped to the patient",
    "The patient picks up (with ID verification and signature) or receives delivery",
]:
    doc.add_paragraph(s, style='List Number')

# The Systems
green_heading("The Systems, Explained")

green_heading("1. The Database (Supabase + Prisma)", level=2)
doc.add_paragraph("What it is: A database is where all the pharmacy's information is stored \u2014 patient names, prescriptions, inventory counts, insurance details, everything. Ours is a PostgreSQL database hosted by Supabase.")
doc.add_paragraph("What Prisma does: Prisma is a translator. Our code is written in TypeScript. Prisma converts our code's requests into database language so the two can talk.")
doc.add_paragraph("What Supabase adds: Login/password verification, real-time updates (when someone changes a record, everyone sees it immediately), and security rules controlling who can see what.")

green_heading("2. The Website Host (Vercel)", level=2)
doc.add_paragraph("What it is: Vercel puts our PMS on the internet. When you open the PMS in a browser, Vercel delivers those pages. It also runs scheduled jobs like the DRX sync every 5 minutes.")

green_heading("3. Insurance Claims (Change Healthcare)", level=2)
doc.add_paragraph("What it is: When a patient has insurance, we ask: 'Do you cover this drug? What's the copay?' Change Healthcare runs the network that carries this question to the insurance company and brings back the answer. The format is NCPDP D.0 \u2014 a nationwide standard.")
doc.add_paragraph("Status: Code is built and tested. Needs vendor contract and certification to go live with real claims.")

green_heading("4. E-Prescribing (SureScripts)", level=2)
doc.add_paragraph("What it is: Doctors send prescriptions electronically. SureScripts is the network that delivers them. Handles new prescriptions, change requests, renewals, and cancellations.")
doc.add_paragraph("Status: Code is built. Needs SureScripts partner credentials. Until then, prescriptions come via our prescriber portal, manual entry, phone, and fax.")

green_heading("5. Automation (Keragon)", level=2)
doc.add_paragraph("What it is: Keragon makes things happen automatically \u2014 sending texts when Rx is ready, creating alerts for low stock, logging events. It runs 24/7 without getting tired or forgetting.")

green_heading("6. The Old System (DRX)", level=2)
doc.add_paragraph("What it is: The pharmacy system we used before. We sync data FROM DRX every 5 minutes. Our PMS can now process prescriptions through all queue steps independently \u2014 there's a toggle to disable DRX sync when ready.")

green_heading("7. Error Monitoring (Sentry)", level=2)
doc.add_paragraph("What it is: Watches the entire application. When something breaks, it captures a detailed report and alerts us. We've added PHI scrubbing so patient data is stripped from error reports.")

green_heading("8. AI Assistant (Claude AI)", level=2)
doc.add_paragraph("What it is: AI built into the PMS to help pharmacists. Checks drug interactions, suggests compounding formulas, drafts counseling notes, and helps with insurance rejections. It's a second opinion tool, not a replacement.")

doc.add_page_break()

# NEW: What We've Built
green_heading("Everything We've Built")
doc.add_paragraph("Here's the full feature list \u2014 this goes way beyond what DRX offers.")

green_heading("Queue Workflow & Fill Processing", level=2)
doc.add_paragraph("Every prescription moves through: Intake \u2192 Adjudicating \u2192 Print \u2192 Scan \u2192 Verify \u2192 Waiting Bin \u2192 Sold. The dashboard shows live counts for all 16 queues. The Fill Processing Screen shows patient info, drug details, allergy warnings, and a panel specific to each step (barcode scanner for Scan, pharmacist checklist for Verify, etc.).")

green_heading("Hardware Integrations", level=2)
for item in [
    "Pharmacy Scales (Sartorius/Ohaus) \u2014 weights read directly into batch records via USB",
    "Eyecon Pill Counter \u2014 99.99% accuracy with NDC verification",
    "Stripe Payment Terminals \u2014 tap/swipe/insert, FSA/HSA support",
    "Zebra Label Printers \u2014 Rx labels, compound labels, batch records",
]:
    doc.add_paragraph(item, style='List Bullet')

green_heading("IVR Phone System", level=2)
doc.add_paragraph("Patients call and hear: Press 1 for refills, Press 2 for status, Press 3 for pharmacist. Creates refill requests automatically and sends SMS confirmation. Also supports automated outbound calling campaigns.")

green_heading("Inventory Management", level=2)
doc.add_paragraph("Optimization Engine: Analyzes dispensing velocity, calculates optimal reorder points and EOQ, detects dead stock, identifies fast movers. Cardinal Health Ordering: Search catalog, create purchase orders, track deliveries \u2014 all in the PMS.")

green_heading("Compounding", level=2)
doc.add_paragraph("Formulas with versions, ingredients, and procedures. Batch creation with ingredient weighing from specific lots (FIFO), QA checks, compounder sign-off, and pharmacist verification. Batch record PDFs for compliance.")

green_heading("Clinical Decision Support", level=2)
for item in [
    "DUR Engine \u2014 drug-drug interactions, allergy cross-reactivity, therapeutic duplication, dose range checks",
    "Immunization Registry \u2014 Louisiana LINKS integration, CDC schedule recommendations, vaccine tracking",
    "PSE Tracking \u2014 pseudoephedrine purchase log with daily/monthly limits (CMEA compliance)",
    "Controlled Substance Perpetual Inventory \u2014 running balance per Schedule II-V drug, biennial report",
]:
    doc.add_paragraph(item, style='List Bullet')

green_heading("Billing & Claims", level=2)
doc.add_paragraph("NCPDP D.0 claim building, rejection resolution with code explanations, charge accounts for clinic billing, AR aging reports, POS with register sessions, and FSA/HSA IIAS compliance.")

green_heading("Shipping & Delivery", level=2)
doc.add_paragraph("Shipment creation with carrier selection, driver route optimization, remote delivery signature capture (mobile page with GPS + photo), compliance packaging (blister packs) with sync list and change tracking.")

green_heading("Analytics Dashboard", level=2)
doc.add_paragraph("KPI cards, dispensing trends, revenue by payer type, top drugs, claims performance, payer mix, productivity metrics, compounding stats. All with 7-day/30-day/90-day/YTD date ranges.")

green_heading("Telepharmacy", level=2)
doc.add_paragraph("Video consultations for remote pharmacist verification and patient counseling. OBRA-compliant counseling checklist. Enables serving remote locations with just a tech on-site.")

green_heading("Web Refill Widget", level=2)
doc.add_paragraph("Embeddable iframe for pharmacy websites. Patients enter name, DOB, and RX number to request refills or check status. API key authentication and rate limiting.")

green_heading("Security & Compliance", level=2)
for item in [
    "All endpoints require authentication",
    "Twilio webhook signature verification (prevents forged IVR requests)",
    "PHI access audit logging (every patient/prescription view logged)",
    "CSP security headers, HSTS, rate limiting",
    "Role-based access control with 2FA for admin/pharmacist",
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Status Table
green_heading("Where Everything Stands Today")
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 3'
table.rows[0].cells[0].text = 'System'
table.rows[0].cells[1].text = 'Status'
table.rows[0].cells[2].text = 'Notes'
for system, status, notes in [
    ("Database (Supabase)", "LIVE", "168K+ patients synced"),
    ("Website (Vercel)", "LIVE", "bndsrxportal.com"),
    ("Queue Workflow", "LIVE", "Local processing, no DRX needed"),
    ("Fill Processing", "LIVE", "Barcode scan + RPh verification"),
    ("Labels", "LIVE", "DRX templates + custom editor"),
    ("Compounding", "LIVE", "Full formula/batch/QA workflow"),
    ("Insurance Claims", "BUILT", "Needs Change Healthcare contract"),
    ("E-Prescribing", "BUILT", "Needs SureScripts certification"),
    ("PDMP", "BUILT", "Needs Louisiana PMP enrollment"),
    ("Scales/Eyecon", "BUILT", "Connect hardware to go live"),
    ("Stripe POS", "BUILT", "Set API key to go live"),
    ("IVR Phone", "BUILT", "Set Twilio env vars"),
    ("Inventory Optimization", "LIVE", "Uses existing data"),
    ("Cardinal Ordering", "BUILT", "Set API key"),
    ("Analytics", "LIVE", "9 metric categories"),
    ("Controlled Substances", "LIVE", "Perpetual inventory"),
    ("PSE Tracking", "LIVE", "CMEA compliance"),
    ("DUR Engine", "LIVE", "19 interactions + class matching"),
    ("Immunization Registry", "BUILT", "Set env vars"),
    ("Compliance Packaging", "LIVE", "Blister pack sync list"),
    ("Return to Stock", "LIVE", "Auto-RTS with claim reversal"),
    ("Route Optimization", "LIVE", "Delivery routing"),
    ("Delivery Signature", "LIVE", "Mobile GPS + photo"),
    ("Telepharmacy", "BUILT", "Set video API key"),
    ("Web Widget", "LIVE", "Embed code in Settings"),
    ("Outbound Calling", "BUILT", "Set Twilio env vars"),
    ("Sentry Monitoring", "LIVE", "PHI scrubbing active"),
]:
    row = table.add_row().cells
    row[0].text = system
    row[1].text = status
    row[2].text = notes

# Blocking
green_heading("What's Blocking Full Go-Live?")
doc.add_paragraph("The PMS can operate independently for most workflows. Three vendor contracts are needed:")
for item in [
    "Change Healthcare \u2014 Real insurance claims (biggest blocker, outreach doc ready)",
    "SureScripts \u2014 Electronic prescriptions from doctors (code built, needs credentials)",
    "Bamboo Health \u2014 PDMP for controlled substances (code built, needs LA enrollment)",
]:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph("All three can run in parallel. Realistic timeline: 10-12 weeks.")

# Glossary
doc.add_page_break()
green_heading("Glossary")
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 3'
table.rows[0].cells[0].text = 'Term'
table.rows[0].cells[1].text = 'What It Means'
for term, defn in [
    ("API", "How two software systems talk to each other"),
    ("CMEA", "Law requiring pseudoephedrine purchase tracking"),
    ("DUR", "Drug Utilization Review \u2014 clinical safety checks"),
    ("EOQ", "Economic Order Quantity \u2014 optimal order size"),
    ("EPCS", "Electronic Prescribing for Controlled Substances"),
    ("FSA/HSA", "Tax-advantaged medical spending accounts"),
    ("HIPAA", "Law protecting patient data privacy"),
    ("IVR", "Automated phone menu ('Press 1 for refills')"),
    ("MRN", "Medical Record Number (ours: BNDS-XXXXXXX)"),
    ("NCPDP D.0", "Standard format for insurance claims"),
    ("NDC", "10-digit code identifying every drug product"),
    ("NPI", "10-digit number identifying healthcare providers"),
    ("PBM", "Company that processes insurance claims"),
    ("PDMP", "State database tracking controlled substances"),
    ("PHI", "Protected Health Information (HIPAA-covered data)"),
    ("PSE", "Pseudoephedrine \u2014 tracked due to meth risk"),
    ("RPh", "Registered Pharmacist"),
    ("RTS", "Return to Stock \u2014 reversing unclaimed prescriptions"),
    ("SIG", "Prescription directions ('Take 1 tablet daily')"),
]:
    row = table.add_row().cells
    row[0].text = term
    row[1].text = defn

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\nQuestions? Ask Caleb (cdowling@bndsrx.com)")
run.font.size = Pt(10)
run.font.color.rgb = light

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "PMS_for_Dummies.docx")
doc.save(out)
print(f"Saved: {out}")
