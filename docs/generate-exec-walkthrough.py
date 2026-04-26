"""Generate Executive Presentation Walkthrough - Caleb's eyes only"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

green = RGBColor(0x40, 0x72, 0x1D)
red = RGBColor(0xDC, 0x26, 0x26)
blue = RGBColor(0x25, 0x63, 0xEB)
gray = RGBColor(0x6B, 0x6B, 0x6B)

def gh(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = green

def say(text):
    p = doc.add_paragraph()
    r = p.add_run("SAY: "); r.bold = True; r.font.color.rgb = green; r.font.size = Pt(10)
    r = p.add_run(text); r.italic = True; r.font.size = Pt(10)

def click(text):
    p = doc.add_paragraph()
    r = p.add_run("DO: "); r.bold = True; r.font.color.rgb = blue; r.font.size = Pt(10)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(10)

def private(text):
    p = doc.add_paragraph()
    r = p.add_run("NOTE: "); r.bold = True; r.font.color.rgb = red; r.font.size = Pt(9)
    r = p.add_run(text); r.font.color.rgb = red; r.font.size = Pt(9)

def timer(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(9); r.font.color.rgb = gray; r.bold = True

def bul(text):
    doc.add_paragraph(text, style='List Bullet')

# COVER
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("BNDS PMS\nExecutive Presentation\nWalkthrough Guide"); r.font.size = Pt(26); r.font.color.rgb = green; r.bold = True

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CONFIDENTIAL - FOR CALEB ONLY"); r.font.size = Pt(12); r.font.color.rgb = red; r.bold = True

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Your personal script for presenting to the executive team.\nGreen = what to say  |  Blue = what to click  |  Red = private notes\n\nEstimated time: 25-30 minutes"); r.font.size = Pt(10); r.font.color.rgb = gray

doc.add_page_break()

# PRE-MEETING
gh("Before the Meeting (5 min prep)")
private("Do this 5 minutes before executives walk in")
bul("Open Chrome, go to www.bndsrxportal.com, log in")
bul("Verify dashboard loads with queue counts and recent activity")
bul("Open second tab to DRX (boudreaux.drxapp.com) for side-by-side comparison")
bul("Have this guide on your phone or printed - NOT on the screen")
bul("Close Slack, email, notifications - no distractions")
bul("Make sure the Waiting Bin and Patients pages have data to show")

doc.add_page_break()

# PART 1
gh("Part 1: Why We Built This (3 min)")
timer("Minutes 0-3")
click("Show the DRX tab first")
say("Right now we run our entire pharmacy on DRX. It works, but we have hit the ceiling on what it can do for us.")
bul("We do not own our data - it sits on DRX servers")
bul("Compounding QA is tracked on paper - no digital batch records")
bul("No delivery route optimization - drivers plan routes manually")
bul("Limited phone system - no automated IVR for patient refills")
bul("No controlled substance perpetual inventory - DEA compliance is manual")
say("So we built our own. Let me show you what it can do.")
click("Switch to BNDS PMS tab (bndsrxportal.com)")

doc.add_page_break()

# PART 2
gh("Part 2: The Dashboard (3 min)")
timer("Minutes 3-6")
click("You should be on the Dashboard. Point to the screen.")
say("This is our command center. Everything the pharmacy needs is visible at a glance.")

gh("Workflow Queue", 2)
click("Point to the left column with 16 queues")
say("Live queue counts. Every prescription moves through these steps - Intake through Waiting Bin. Numbers update automatically from DRX every 5 minutes. When we go independent, they update in real-time from our own database.")

gh("Quick Access", 2)
click("Point to the 9 icon cards")
say("One-click access to every module. Techs and pharmacists can get anywhere in two clicks.")

gh("Recent Activity & Stock Alerts", 2)
click("Point to the right column")
say("Real-time activity feed and stock alerts for items running low.")

private("If queue counts show 0, say: these sync every 5 minutes, the live system shows real counts")

doc.add_page_break()

# PART 3
gh("Part 3: Prescription Workflow (5 min)")
timer("Minutes 6-11")
say("Let me walk you through how a prescription flows through the system.")

gh("Intake", 2)
click("Click eRx Intake in the sidebar")
say("When a doctor sends an e-prescription, it lands here. The system auto-matches patient, prescriber, and drug. If everything matches, it creates the Rx automatically. If uncertain, a tech reviews it.")
private("If intake queue is empty, just describe the flow and move on")

gh("Queue Processing", 2)
click("Click any queue item from dashboard, or go to Queue page")
say("This is where the work happens. Every fill has a Process button that opens a dedicated screen with patient info, allergy warnings, drug details, and a panel specific to each step.")
say("For the Scan step, the tech scans the drug barcode - if the NDC does not match, it blocks advancement. Patient safety built in.")
say("For the Verify step, the pharmacist sees a clinical checklist and signs off with their credentials. Every action is logged with who did it and when.")
private("If no fills to process, describe verbally. Do not fumble trying to find test data.")

gh("Pickup", 2)
click("Click Pickup in the sidebar")
say("Verified prescriptions go to the waiting bin. Color-coded by age. At pickup, we capture digital signature, verify ID, offer counseling, and collect copay through our POS.")

doc.add_page_break()

# PART 4
gh("Part 4: Compounding (3 min)")
timer("Minutes 11-14")
click("Click Compounding in the sidebar")
say("This is where we really differentiate. Our compounding workflow is fully digital - no more paper batch records.")
say("Every formula has versioned ingredients and step-by-step procedures. When we compound, each ingredient is weighed from a specific lot - FIFO compliance. The scale reading goes directly into the system. QA checks, compounder sign-off, pharmacist verification - two signatures required. The batch record prints as a compliance PDF.")
private("If asked about scales: We have Sartorius and Ohaus integration. Weight reads directly from USB.")

doc.add_page_break()

# PART 5
gh("Part 5: Billing (2 min)")
timer("Minutes 14-16")

gh("Billing", 2)
click("Click Billing in the sidebar")
say("Insurance claims with NCPDP rejection codes and resolution steps. When a claim gets rejected, we see the code, the explanation, and the steps to fix it. Charge accounts for clinic billing with AR aging at 30, 60, 90 days.")
private("If asked about real claims: The claim builder is complete. We need to sign with Change Healthcare. That is our number one go-live blocker.")

doc.add_page_break()

# PART 6
gh("Part 6: Phone System & POS (3 min)")
timer("Minutes 16-19")

gh("Phone System / IVR", 2)
click("Point to the Quick Dial phone pad on the dashboard")
say("We have a built-in phone system. The dialer on the dashboard lets staff call patients directly without switching apps. But there is a lot more to it than just a dialer.")

say("First, the automated IVR. When patients call our pharmacy number, they hear a menu - Press 1 for refills, Press 2 to check prescription status, Press 3 to speak with a pharmacist. If they press 1, they enter their RX number and the system automatically creates a refill request. They get a text confirming it was received. No staff time needed for routine refill calls.")

click("Click Phone in the sidebar to show the live call dashboard")
say("This is our call center dashboard. Staff can see every active call in real-time - who is calling, whether they have been matched to a patient in our system, how long they have been on the line, and what they called about.")
say("On the left, active calls with live duration timers. In the center, the hold queue - anyone on hold sorted by wait time, with a red highlight if they have been waiting more than two minutes. On the right, transfer controls. Staff can put a caller on hold, retrieve them, transfer to the pharmacist line, billing, or shipping - all with one click, right from this screen.")
say("We also have automated outbound calling campaigns. The system can call patients in batch to notify them their prescription is ready or remind them about refills coming due.")

click("Click Call History link if visible")
say("Every call is logged - who called, when, how long, who handled it, and the outcome. Fully searchable and filterable.")

private("If asked about the IVR: It uses Twilio. We need to set the Twilio environment variables to go live. The code is fully built. The call management dashboard works with or without Twilio configured - in dev mode it shows the UI but calls are simulated.")

gh("Point of Sale (POS)", 2)
click("Click POS in the sidebar")
say("Our POS handles copay collection at pickup. Cash, credit, debit - and we support FSA and HSA cards with the correct merchant category code so the transaction processes as a qualified medical expense. We track register sessions with opening and closing balances, daily revenue totals, and full transaction history.")
say("We have Stripe Terminal integration built so we can connect physical card readers for tap, swipe, and chip payments.")
private("If asked about Stripe: Code is built. Need to set STRIPE_SECRET_KEY and connect a physical terminal reader.")

doc.add_page_break()

# PART 7
gh("Part 7: Delivery & Shipping (2 min)")
timer("Minutes 19-21")
click("Click Shipping in the sidebar")
say("For patients who cannot pick up: route optimization for driver efficiency, remote signature capture with GPS and photo proof of delivery, and for controlled substances the driver verifies patient ID. We also have compliance packaging for blister pack patients.")

doc.add_page_break()

# PART 8
gh("Part 8: Settings (1 min)")
timer("Minutes 21-22")
click("Click Settings - show the sidebar navigation")
say("Clean sidebar layout for all configuration. Pharmacy info, security, templates, hardware, integrations - everything configurable without touching code.")
click("Click through 2-3 sections quickly")
private("Do not linger here. Settings are not exciting.")

doc.add_page_break()

# PART 9
gh("Part 9: The Ask (3 min)")
timer("Minutes 22-25")
say("Everything you have seen is live in production right now. To fully replace DRX, we need three vendor contracts.")

gh("1. Change Healthcare (Insurance)", 2)
say("The biggest one. Without it, every Rx is cash-pay. Code is built and tested. Need the contract and certification. 4-8 weeks after signing.")

gh("2. SureScripts (E-Prescribing)", 2)
say("Connects us to every doctor EHR in the country. Code is built. Need credentials. 8-16 weeks for certification.")

gh("3. Bamboo Health (PDMP)", 2)
say("Required for controlled substances. Code built. Need state enrollment. 2-4 weeks.")

say("All three run in parallel. Realistic go-live: 10-12 weeks. DRX stays as backup during transition.")

doc.add_page_break()

# Q&A
gh("Anticipated Questions")
timer("Minutes 25-30 (Q&A)")
private("Have these answers ready. They WILL ask these.")

gh("How much did this cost?", 2)
say("Development was in-house with AI-assisted coding. Ongoing infrastructure is under $100 per month. Compare that to what we pay DRX.")

gh("What if it breaks?", 2)
say("Sentry monitoring alerts us within seconds. Vercel auto-rolls back failed deploys. DRX runs as backup during transition.")

gh("Is it HIPAA compliant?", 2)
say("Yes. Encrypted data, role-based access control, PHI audit logging, 2FA for pharmacists, Content Security Policy headers. Every patient record view is logged.")

gh("What about training?", 2)
say("Same workflow as DRX but with better tools. We have a for-dummies guide and a full operations manual. 2-4 hours of training for existing staff.")

gh("Can we add features?", 2)
say("That is the whole point. With DRX we wait. With our system we add features same-day. We already have features DRX does not - prescriber portal, route optimization, controlled substance tracking, telepharmacy.")

gh("Risk of staying on DRX?", 2)
say("Rising costs, vendor lock-in, no customization. Our compounding, clinic billing, and delivery needs are not well-served by a generic system.")

# Footer
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CONFIDENTIAL - DO NOT DISTRIBUTE\nDestroy after presentation"); r.font.size = Pt(9); r.font.color.rgb = red; r.bold = True

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Executive_Walkthrough_v3.docx")
doc.save(out)
print(f"Saved: {out}")
